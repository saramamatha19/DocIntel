import io
from pathlib import Path
import pymupdf
import pytesseract
import docx
from PIL import Image
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse

def extract_document(file_path: str) -> list[dict]:
    """
    Extract content from a document, dispatching to the right
    extractor based on file extension. All extractors return
    the same normalized list-of-dicts shape, so nothing else in
    the pipeline (chunking, embedding, classification) needs to
    know or care what format a document originally was.
    """

    extension = Path(file_path).suffix.lower()

    if extension == ".pdf":
        return extract_pdf(file_path)

    if extension == ".docx":
        return extract_docx(file_path)

    if extension in (".txt", ".md"):
        return extract_plain_text(file_path)

    raise ValueError(
        f"Unsupported file type: {extension}"
    )


def extract_pdf(pdf_file: str) -> list[dict]:
    """
    Extract text, tables, and OCR text from images in a PDF.
    Returns a normalized list of content items.
    """
    document = pymupdf.open(pdf_file)
    document_name = Path(pdf_file).name
    content = []

    for page_number, page in enumerate(document, start=1):
        # 1. Detect tables
        tables = page.find_tables()
        table_rects = [
            pymupdf.Rect(table.bbox)
            for table in tables.tables
        ]

        # 2. Extract normal text
        text_blocks = []
        for block in page.get_text("blocks"):

            block_rect = pymupdf.Rect(block[:4])
            block_text = block[4]

            # Don't include text that belongs to a table.
            inside_table = any(
                block_rect.intersects(table_rect)
                for table_rect in table_rects
            )

            if not inside_table and block_text.strip():
                text_blocks.append(block_text.strip())

        normal_text = "\n".join(text_blocks)

        if normal_text.strip():
            content.append({
                "text": normal_text,
                "document_name": document_name,
                "page_number": page_number,
                "content_type": "text",
            })

        # 3. Extract tables
        for table_index, table in enumerate(
            tables.tables,
            start=1
        ):

            try:
                table_data = table.extract()

                # Convert table rows into readable text.
                table_lines = []

                for row in table_data:

                    cells = [
                        str(cell).strip() if cell is not None else ""
                        for cell in row
                    ]

                    # A row that's really a caption/footnote (real
                    # content only in the first cell) otherwise
                    # leaves a dangling "|  |" from its empty
                    # trailing cells once joined.
                    while cells and not cells[-1]:
                        cells.pop()

                    row_text = " | ".join(cells)

                    table_lines.append(row_text)

                table_text = "\n".join(table_lines)

                if table_text.strip():
                    content.append({
                        "text": table_text,
                        "document_name": document_name,
                        "page_number": page_number,
                        "content_type": "table",
                        "table_number": table_index,
                    })

            except Exception as exc:
                print(
                    f"Warning: Could not extract table "
                    f"{table_index} on page {page_number}: {exc}"
                )

         # 4. Extract images and OCR their text
        image_list = page.get_images(full=True)
        for image_index, image_info in enumerate(
            image_list,
            start=1
        ):

            try:
                xref = image_info[0]

                base_image = document.extract_image(xref)

                image_bytes = base_image["image"]

                image = Image.open(
                    io.BytesIO(image_bytes)
                )

                ocr_text = pytesseract.image_to_string(
                    image
                ).strip()

                if ocr_text:
                    content.append({
                        "text": ocr_text,
                        "document_name": document_name,
                        "page_number": page_number,
                        "content_type": "image_ocr",
                        "image_number": image_index,
                    })

            except Exception as exc:
                print(
                    f"Warning: Could not OCR image "
                    f"{image_index} on page {page_number}: {exc}"
                )

    document.close()

    return content


def extract_plain_text(file_path: str) -> list[dict]:
    """
    Extract content from a .txt or .md file. Neither format
    needs a parsing library — Markdown syntax is left as-is in
    the text rather than stripped, same as how PDF extraction
    doesn't try to interpret formatting beyond raw text.

    There's no real concept of "pages" in a plain text file, so
    the whole file is treated as a single page.
    """

    document_name = Path(file_path).name

    text = Path(file_path).read_text(
        encoding="utf-8",
        errors="ignore",
    )

    if not text.strip():
        return []

    return [{
        "text": text.strip(),
        "document_name": document_name,
        "page_number": 1,
        "content_type": "text",
    }]


def extract_docx(file_path: str) -> list[dict]:
    """
    Extract content from a Word (.docx) file by joining its
    paragraphs into one block of text. Word doesn't expose page
    breaks in a way that's reliably extractable without actually
    rendering the document, so — same as plain text — the whole
    file is treated as a single page.
    """

    document_name = Path(file_path).name

    word_document = docx.Document(file_path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in word_document.paragraphs
        if paragraph.text.strip()
    ]

    text = "\n".join(paragraphs)

    if not text.strip():
        return []

    return [{
        "text": text,
        "document_name": document_name,
        "page_number": 1,
        "content_type": "text",
    }]


#Extract text from a URL
def extract_webpage(url: str) -> list[dict]:
    """
    Extract readable text from a webpage.

    Returns the same normalized structure used by PDF extraction.
    """

    response = requests.get(
        url,
        timeout=15,
        headers={
            "User-Agent": "DocIntel/1.0"
        },
    )

    response.raise_for_status()

    soup = BeautifulSoup(
        response.text,
        "html.parser",
    )

    # Remove elements that usually don't contain
    # useful document content.
    for element in soup([
        "script",
        "style",
        "nav",
        "footer",
        "header",
    ]):
        element.decompose()

    text = soup.get_text(
        separator="\n",
        strip=True,
    )

    # Clean empty lines
    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    clean_text = "\n".join(lines)

    parsed_url = urlparse(url)

    document_name = (
        parsed_url.netloc
        + parsed_url.path
    )

    return [{
        "text": clean_text,
        "document_name": document_name,
        "page_number": 1,
        "content_type": "webpage",
        "url": url,
    }]



# Test --pdf text
'''if __name__ == "__main__":

    pdf_file = "uploads/atlassian/Security_Measures.pdf"

    content = extract_document(pdf_file)

    for index, item in enumerate(content, start=1):

        print(
            f"\n--- Item {index} | "
            f"{item['content_type'].upper()} | "
            f"Page {item['page_number']} ---"
        )

        print(f"Document: {item['document_name']}")

        print(item["text"])'''

if __name__ == "__main__":

    url = "https://www.atlassian.com/software"

    content = extract_webpage(url)

    for index, item in enumerate(
        content,
        start=1,
    ):

        print(
            f"\n--- Item {index} | "
            f"{item['content_type'].upper()} ---"
        )

        print(
            f"Document: {item['document_name']}"
        )

        print(
            f"URL: {item['url']}"
        )

        print(
            item["text"][:2000]
        )